#!/usr/bin/env python3
"""Build the July 21 checkpoint report as an editable .docx.

Word styles (Title / Heading 1 / List Bullet) rather than hand-formatted
runs, so the file stays editable in Word or LibreOffice. Text lives in
BODY below; edit there and re-run.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

TITLE = ("NL2Plan: natural-language task planning on the mobile "
         "manipulator — checkpoint report")
BYLINE = "Rajeev Reddy — July 21, 2026"

DIAGRAM = """\
  "pick up the magenta block and place it on the table"
                        |
                        v
 +--------------------------+  tool call   +---------------------+
 |  agent loop              |------------->|  tool dispatcher    |
 |  Qwen2.5-7B-Instruct     |<-------------|  + ROS2 backend     |
 |  (Ollama, local)         |  result      +---------------------+
 |  step/time caps, schema  |                |      |         |
 |  checks, JSONL trace     |                v      v         v
 +--------------------------+              Nav2  stop-spin  arm and
                                          goals  scan and   gripper
                                                 confirm    control
                                                    ^
                                   /block_pose/<color> (map frame)
                                   multi-colour HSV detector:
                                   per-colour bands, every candidate
                                   blob, size-vs-distance gate"""

# (heading, [blocks]) where a block is ("p", text), ("b", (lead, rest)),
# ("i", text) for an intro paragraph, or ("mono", text).
BODY = [
    ("What it does", [
        ("p", "At the last checkpoint the robot ran one mission, and a state "
              "machine I wrote by hand decided every step of it. That state "
              "machine is gone. The robot now takes typed commands in plain "
              "English, and a 7-billion-parameter language model running "
              "locally on my machine decides which capabilities to invoke, in "
              "what order, and with what arguments."),
        ("p", "The model has four tools: navigate_to, find_object, pick and "
              "place. It calls them against the same Nav2 stack, camera "
              "detector and arm controllers as the last checkpoint, reads the "
              "result of every call, and replans when one fails. An empty scan "
              "sends it to another room and back. A refused grasp sends it "
              "back to re-verify the block before trying again. Commands run "
              "inside a session that keeps its history, so after fetching a "
              "few blocks I can ask “where did you find each block?” "
              "and get an answer from memory, with the robot standing still."),
        ("p", "The house now holds four 5 cm blocks (red, orange, magenta, "
              "brown), one per area, each roughly 0.7 m from a named search "
              "pose. On the final build, twelve missions ran back to back "
              "without a failure: one recorded session that fetched all four "
              "blocks in a single conversation at six model steps each, and a "
              "scripted rehearsal that ran every block twice from a freshly "
              "reset scene. The rehearsal logged no failed tool calls at all, "
              "with missions taking four to eight steps and a little over two "
              "minutes each."),
    ]),
    ("Architecture", [
        ("p", "The division of labour follows the SayCan line of work: the "
              "language model does task-level reasoning and nothing else, and "
              "every capability underneath it is ordinary, verifiable "
              "robotics."),
        ("mono", DIAGRAM),
        ("b", ("Planner.", " Qwen2.5-7B-Instruct through Ollama's tool-use "
               "API, running offline. A small local model is a deliberate "
               "choice: how reliable a 7B model can be made under tool use is "
               "part of what I wanted to find out. The loop around it enforces "
               "step and wall-clock caps, validates arguments against JSON "
               "schemas before dispatch, prods the model when it narrates a "
               "plan instead of calling a tool, and writes every call and "
               "result to a JSONL trace. Nearly all the debugging described "
               "below came out of that trace.")),
        ("b", ("The prompt as a semantic map.", " The system prompt names the "
               "area each block lives in, plus a fixed order to sweep if a "
               "scan comes up empty. It only decides which room to drive to "
               "first; the perception work still happens locally and honestly "
               "once the robot gets there.")),
        ("b", ("Perception.", " The single-target detector became a "
               "multi-colour node with one pose topic per colour. Orange and "
               "brown turned out to share hue almost exactly, so saturation is "
               "what separates them: orange renders fully saturated, brown "
               "sits near 197, wooden furniture near 152. Every blob above the "
               "area threshold is offered as a candidate rather than just the "
               "largest one, because furniture regularly out-sizes a 5 cm "
               "cube, and the size-versus-distance gate decides which "
               "candidate is really a block.")),
        ("b", ("The confirm gate.", " Nothing counts as found until the robot "
               "has stopped and four consecutive sightings cluster within "
               "0.2 m of each other in the robot’s own frame, at a "
               "distance a real floor block could plausibly occupy. Both of "
               "those details cost me time to get right and are covered "
               "below.")),
        ("b", ("Grasp approach.", " Nav2 plans the approach to a 0.6 m "
               "standoff, since a blind creep from the search pose once drove "
               "into a stool. The block is then re-confirmed dead ahead, where "
               "the projection error largely disappears, and only the last "
               "stretch runs on odometry. If the block cannot be re-confirmed, "
               "or the creep presses into something before reaching grasp "
               "range, the pick is refused and the model is told to scan "
               "again. Grasping itself is still the pinned attachment from the "
               "last checkpoint, and placement is still the touch-dock.")),
    ]),
    ("What I found", [
        ("b", ("A 7B planner needs the world written down.", " Without the "
               "block-to-room list in the prompt, the model wandered from room "
               "to room and hit the step cap before it ever reached a pick. "
               "With it, the first navigation is always right and missions "
               "settle at four to eight steps. The wording matters as much as "
               "the content: an earlier version of the search rule happened to "
               "name the lounge first, and the model went to the lounge for "
               "every block regardless of what I asked for.")),
        ("b", ("Error messages are the model’s interface.", " A small "
               "model recovers about as well as the error strings let it. When "
               "a malformed pick only got back “run find_object "
               "first”, the model re-scanned and made the same mistake "
               "again until it ran out of steps. Once the error named the "
               "identifier that would have worked, it corrected on the next "
               "call. I ended up treating tool errors as part of the API "
               "design rather than as diagnostics for me.")),
        ("b", ("Cluster in the robot’s frame, not the map.", " At one "
               "search pose AMCL believed a heading 24° away from ground "
               "truth. That error is baked into every pixel the detector "
               "back-projects, so map-frame sightings of a stationary block "
               "smear well past any sensible cluster threshold and the robot "
               "reports “not visible” while staring straight at it. "
               "Measured relative to the robot the same sightings are steady, "
               "because the identical error corrupts the robot pose and the "
               "projection and cancels out of the difference. The confirm gate "
               "now clusters robot-relative, and the final approach re-anchors "
               "the sighting through whatever the robot currently believes its "
               "pose to be.")),
        ("b", ("Localization quality is a property of place.", " That 24° "
               "error was not drift I could tune away; it belonged to one "
               "corner of the house with nothing distinctive for the lidar to "
               "match against. I proved it by swapping two blocks: the failure "
               "stayed with the corner, not with the colour. Moving that block "
               "to the floor behind the living-room sofa, where a large flat "
               "surface sits in view, brought the heading error at the search "
               "pose down to about 3.5° and the failure disappeared "
               "without touching AMCL at all.")),
        ("b", ("Gates have to match where the robot actually stops.", " The "
               "confirm gate’s distance ceiling came from the nominal "
               "search poses, which sit 0.7 to 0.8 m from each block. But Nav2 "
               "is allowed a quarter metre of goal tolerance, and with "
               "localization error on top the robot routinely parks a little "
               "over a metre away, so the gate was throwing away good "
               "sightings. On screen this looked exactly like a detector bug: "
               "drive to the block, detect it, sit there, leave. What settled "
               "it was making every failed scan report the reason each "
               "candidate was rejected. The log then showed clean clusters "
               "sitting on the block’s true position being discarded at "
               "0.95 to 1.19 m, and the fix was a single measured constant.")),
        ("b", ("Never compare two localization fixes seconds apart to make a "
               "sub-metre decision.", " I added a safety check that "
               "re-measured the robot-to-block distance after the final creep, "
               "and it started refusing picks where the robot was plainly "
               "standing at the block. AMCL had simply corrected itself during "
               "the creep, and the check was reading that correction as half a "
               "metre of shortfall. Short-horizon geometry belongs to "
               "odometry. Moving the same check inside the creep, where it is "
               "measured the way the robot actually drives, made the false "
               "refusals disappear.")),
        ("b", ("Refuse rather than pretend.", " A pinned grasp succeeds from "
               "anywhere, which means honesty has to be enforced before the "
               "grasp rather than by it. The old code would act on a stale "
               "sighting and teleport the block across a visible gap into the "
               "gripper, which is worthless in a demo and worse than worthless "
               "in a measurement. Any grasp that cannot be re-verified is now "
               "refused. A refusal costs one extra scan; a fake success "
               "quietly corrupts everything measured after it. Several "
               "apparent regressions during this stage turned out to be old "
               "silent failures finally becoming visible.")),
    ]),
    ("Results", [
        ("p", "All of the following ran at a real-time factor near 1.0 on the "
              "final build."),
        ("b", ("Twelve consecutive successful missions across all four "
               "blocks.", " One recorded session fetched all four in a single "
               "conversation at six model steps each, and a scripted rehearsal "
               "ran every block twice from a freshly reset scene: eight of "
               "eight, with no failed tool calls anywhere in the batch, at "
               "four to eight steps and 132 to 154 seconds per mission.")),
        ("b", ("Recovery exercised for real.", " Earlier builds the same day "
               "produced the failures worth having: empty scans led the robot "
               "to search elsewhere and return, refused grasps led to a "
               "re-scan and then a successful pick, and empty model replies "
               "were nudged back on task mid-mission rather than ending it.")),
        ("b", ("Session memory.", " After a run of fetch commands, asking "
               "where each block was found produces a correct answer in one "
               "step with no robot motion.")),
        ("b", ("36 unit tests pass.", " Each perception and approach lesson "
               "above is pinned by a regression test built from the live "
               "failure it came from, so the same mistake cannot quietly "
               "return.")),
        ("p", "One blemish is visible in the recorded video and worth naming. "
              "Every placement currently aims at the same point on the drop "
              "table, so by the third and fourth block the drop zone is "
              "crowded and two blocks ended up settling beside the table "
              "rather than on it. The missions themselves succeeded; the fix "
              "is to give each block its own drop spot, and it is first on the "
              "list below."),
    ]),
    ("Progress toward reusable, recallable skill modules", [
        ("p", "The suggestion after the last checkpoint was to decompose the "
              "simulation into reusable hybrid modules that can be recalled "
              "from a knowledge base and be sequentially configured and "
              "parameterized to achieve any task outcome, with random "
              "interruptions. That shaped what I built this stage. Taking the "
              "pieces of it one at a time:"),
        ("b", ("Reusable modules: done.", " This checkpoint is that refactor. "
               "The one-mission state machine has been replaced by four "
               "self-contained capabilities behind a single call-and-result "
               "interface. Each carries its own verification and its own "
               "failure semantics, returns a structured result instead of "
               "dying, and can be called in any order with any arguments. The "
               "same four modules served every command anyone typed this "
               "week.")),
        ("b", ("Hybrid: done.", " A learned component chooses and sequences, "
               "while the inside of each module is classical robotics: Nav2 "
               "planning, a geometric detector with explicit gates, scripted "
               "arm motions. Neither half is asked to do the other’s "
               "job.")),
        ("b", ("Sequenced, configured and parameterized at runtime: done.",
               " No sequence is written by hand any more. The planner composes "
               "the modules and fills in their arguments from a "
               "natural-language goal, so a different block, a different room "
               "or a different order needs no new code. Twelve consecutive "
               "missions is my evidence that the composition is dependable "
               "rather than merely possible.")),
        ("b", ("Random interruptions: handled in practice, not yet tested "
               "deliberately.", " Every interruption that arose on its own "
               "this week was absorbed by the replanning loop, including "
               "failed scans, refused grasps and aborted navigation. What I do "
               "not have yet is fault injection I control. The hooks exist "
               "already, since blocks can be teleported mid-mission through "
               "Gazebo’s entity-state service and navigation goals can be "
               "cancelled in flight, so this is a scripting job rather than an "
               "open problem.")),
        ("b", ("Recall from a knowledge base: not yet.", " The skill "
               "vocabulary and the world knowledge still live in the prompt, "
               "as a fixed tool registry plus the map of which block sits in "
               "which area. Nothing is retrieved from an external store. The "
               "decomposition is what makes that step small: each module "
               "already has a name, a typed parameter schema and known "
               "preconditions and effects, so moving those descriptions into a "
               "queryable store and retrieving the relevant subset per command "
               "changes what fills the context window rather than how the "
               "runtime works. The range of achievable tasks would then grow "
               "with the size of the skill base instead of with the length of "
               "the prompt.")),
        ("p", "So the decomposition and the runtime sequencing are in place "
              "and demonstrated, while knowledge-base recall and deliberate "
              "interruption testing are the next stage. I built it so that "
              "both are additions rather than rewrites."),
    ]),
    ("Future work", [
        ("p", "Planned next, roughly in order:"),
        ("b", ("Per-block drop spots.", " Give each block its own place on the "
               "drop table so a later placement stops disturbing an earlier "
               "one. This is the last blemish in the continuous four-block "
               "demo.")),
        ("b", ("Quantitative evaluation.", " Three commands, ten runs each, "
               "reporting success rate, recovery rate and replan latency as a "
               "metrics table in the README.")),
        ("b", ("A skill knowledge base.", " Move the module descriptions, "
               "parameter schemas, preconditions and effects, along with the "
               "scene knowledge, out of the prompt and into a queryable store, "
               "and retrieve only what a command needs.")),
        ("b", ("A scripted interruption suite.", " Deliberate faults rather "
               "than the organic ones seen so far: teleport a block "
               "mid-mission, cancel a navigation goal in flight, block a "
               "doorway during a carry, then measure recovery rate and "
               "time-to-recover for each fault type.")),
        ("b", ("Randomized spawn and a pure search mode.", " Drop the semantic "
               "map from the prompt, randomize block positions per run from a "
               "seed, and let the sweep rule carry the search. It is the "
               "honest robustness test, and it makes “knows the "
               "house” versus “searches the house” a switchable "
               "demo.")),
        ("b", ("Open-vocabulary perception.", " Replace the HSV detector with "
               "GroundingDINO so distractors no longer have to differ by "
               "colour, and a command can name an object the detector was "
               "never tuned for.")),
        ("b", ("Contact-based grasping.", " Carried over from the last "
               "checkpoint: replace the pinned attachment with real finger "
               "contact and friction.")),
    ]),
]

OUT = "/home/reddy/Downloads/checkpoint_report_2.docx"


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in (("Title", 19), ("Heading 1", 13.5)):
        st = doc.styles[style_name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    title = doc.add_paragraph(TITLE, style="Title")
    title.paragraph_format.space_after = Pt(2)

    byline = doc.add_paragraph(BYLINE)
    byline.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    byline.paragraph_format.space_after = Pt(14)

    for heading, blocks in BODY:
        h = doc.add_heading(heading, level=1)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        for kind, content in blocks:
            if kind == "p":
                doc.add_paragraph(content)
            elif kind == "b":
                lead, rest = content
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(6)
                p.add_run(lead).bold = True
                p.add_run(rest)
            elif kind == "mono":
                for line in content.split("\n"):
                    p = doc.add_paragraph()
                    pf = p.paragraph_format
                    pf.space_after = Pt(0)
                    pf.space_before = Pt(0)
                    pf.line_spacing = 1.0
                    pf.left_indent = Pt(18)
                    r = p.add_run(line if line else " ")
                    r.font.name = "Consolas"
                    r.font.size = Pt(7.5)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
